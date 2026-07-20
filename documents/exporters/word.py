"""python-docx DOCX exporters — one template per doc_type."""
import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def _title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)


def _kv(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Light List Accent 1'
    for k, v in rows:
        cells = table.add_row().cells
        cells[0].text = str(k)
        cells[1].text = str(v)


def _narrative(doc, narrative):
    blocks = [b.strip() for b in (narrative or '').split('\n\n') if b.strip()]
    for b in (blocks or ['(no narrative)']):
        doc.add_paragraph(b)


# ── AO-form layout helpers (docs/…/Arrest Warrant Template.pdf, Search
#    Warrant Template.pdf) ──────────────────────────────────────────────
def _form_stamp(doc, text):
    """Small form-identifier line, e.g. 'AO 442 (Rev. 11/11) Arrest Warrant'."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(8)


def _is_federal(court, officer):
    return (court.get('jurisdiction_type_override')
            or officer.get('agency_jurisdiction_type') or 'state') == 'federal'


def _add_seal(doc, officer):
    """Embed the agency seal image; never fail export over it."""
    key = officer.get('agency_seal_key')
    if not key:
        return
    try:
        from utils.storage import get_upload_bytes
        data = get_upload_bytes(key)
        if not data:
            return
        import io as _io
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(_io.BytesIO(data), width=Pt(54))
    except Exception:  # noqa: BLE001 — cosmetic only, never block export
        pass


def _centered(doc, text, bold=False, size=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)


def _court_head(doc, court, officer):
    """Dynamic jurisdiction header (requirement #1) — admin-configured per agency."""
    if _is_federal(court, officer):
        district = court.get('district') or officer.get('agency_judicial_district') \
            or '__________ District of __________'
        _centered(doc, 'UNITED STATES DISTRICT COURT', bold=True, size=14)
        _centered(doc, 'for the', size=10)
        _centered(doc, district, size=11)
        return

    _add_seal(doc, officer)
    state = officer.get('agency_state') or '_______________________'
    county = officer.get('agency_county') or '_____________________'
    court_name = officer.get('agency_court_caption') or officer.get('agency_court_name') \
        or '__________________ COURT'
    _centered(doc, f"STATE OF: {state.upper()}")
    _centered(doc, f"COUNTY OF: {county.upper()}")
    _centered(doc, f"IN THE {court_name.upper()}", bold=True)
    if officer.get('agency_judicial_district'):
        _centered(doc, f"JUDICIAL DISTRICT: {officer['agency_judicial_district'].upper()}")
    if officer.get('agency_division'):
        _centered(doc, f"DIVISION: {officer['agency_division'].upper()}")
    if officer.get('agency_name'):
        _centered(doc, f"AGENCY: {officer['agency_name'].upper()}")


def _caption(doc, left_lines, case_number):
    """The two-column case caption block (parties | Case No.)."""
    table = doc.add_table(rows=1, cols=2)
    left, right = table.rows[0].cells
    for i, line in enumerate(left_lines):
        p = left.paragraphs[0] if i == 0 else left.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(line)
    right.text = f"Case No. {case_number or '__________________'}"
    doc.add_paragraph()


def _checkbox_line(doc, options, selected_key):
    doc.add_paragraph('    '.join(
        f"{'☒' if key == selected_key else '☐'} {label}" for key, label in options))


def _sig_pair(doc, left_label, right_caption):
    """'Date: ___' on the left, a signature line with an italic caption on the right."""
    table = doc.add_table(rows=1, cols=2)
    left, right = table.rows[0].cells
    left.text = f"{left_label} _______________________"
    right.text = '______________________________________'
    cap = right.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(right_caption)
    run.italic = True
    run.font.size = Pt(8)


def _field(doc, label, value=''):
    doc.add_paragraph(f"{label}: {value or '_________________________________'}")


_BANNER_LABELS = {
    'pending_supervisor': 'DRAFT — PENDING SUPERVISOR REVIEW',
    'pending_prosecutor': 'DRAFT — PENDING PROSECUTOR REVIEW',
    'rejected': 'DRAFT — REVIEW REJECTED, NOT APPROVED FOR FILING',
}


def _draft_banner(doc, doc_meta):
    if not doc_meta:
        return
    label = _BANNER_LABELS.get(doc_meta.get('review_status'))
    if not label:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label)
    run.bold = True
    from docx.shared import RGBColor
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)


def _signature(doc, officer, doc_meta=None):
    doc_meta = doc_meta or {}
    signature_name = doc_meta.get('signature_name')
    signed_at = doc_meta.get('signed_at')
    doc.add_paragraph('\n')
    if signature_name and signed_at:
        doc.add_paragraph(f"/s/ {signature_name}")
        doc.add_paragraph(f"{officer.get('full_name', '')}, {officer.get('rank', '')}")
        doc.add_paragraph(f"Badge: {officer.get('badge_number', '')}  ORI: {officer.get('ori', '')}")
        doc.add_paragraph(f"Electronically signed on {signed_at}")
        return
    doc.add_paragraph('_______________________________')
    doc.add_paragraph(f"{officer.get('full_name', '')}, {officer.get('rank', '')}")
    doc.add_paragraph(f"Badge: {officer.get('badge_number', '')}  ORI: {officer.get('ori', '')}")


# ── Incident report ──────────────────────────────────────────────────
def _grid_cell(cell, label, value=''):
    """Template-style cell: tiny label on top, bold value underneath. Empty
    boxes show the label only, exactly like the printed form."""
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(label)
    run.font.size = Pt(6)
    val = str(value).strip() if value is not None else ''
    if val in ('', '-'):
        return
    p2 = cell.add_paragraph()
    run2 = p2.add_run(val)
    run2.bold = True
    run2.font.size = Pt(8)


def _vband(cell, text):
    """Vertical (bottom-to-top) band text, like the form's printed side letters."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches

    tc_pr = cell._tc.get_or_add_tcPr()
    td = OxmlElement('w:textDirection')
    td.set(qn('w:val'), 'btLr')
    tc_pr.append(td)
    cell.width = Inches(0.25)
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(6)


def _band_col(table, text):
    """Merge a table's first column into the narrow letter band."""
    band_cell = table.rows[0].cells[0]
    for r in range(1, len(table.rows)):
        band_cell = band_cell.merge(table.rows[r].cells[0])
    _vband(band_cell, text)
    return band_cell


def _grid(doc, rows_spec, band=None):
    """Bordered grid table; rows_spec is a list of rows of (label, value) or None.
    band adds the template's narrow left letter column ('' keeps it blank)."""
    offset = 0 if band is None else 1
    table = doc.add_table(rows=len(rows_spec), cols=len(rows_spec[0]) + offset)
    table.style = 'Table Grid'
    for r, row in enumerate(rows_spec):
        for c, spec in enumerate(row):
            if spec is not None:
                _grid_cell(table.rows[r].cells[c + offset], spec[0], spec[1])
    if band is not None:
        _band_col(table, band)
    return table


def _legend_row(doc, text, band=None):
    offset = 0 if band is None else 1
    table = doc.add_table(rows=1, cols=1 + offset)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[offset]
    cell.text = ''
    run = cell.paragraphs[0].add_run(text)
    run.font.size = Pt(6)
    if band is not None:
        _vband(table.rows[0].cells[0], band)
    return table


def _incident(doc, form_data, narrative, officer):
    """Smyrna-style INCIDENT/INVESTIGATION REPORT grid (docs/Report Zip)."""
    from docx.shared import Inches

    inc = form_data.get('incident', {})
    facts = form_data.get('facts', {})
    parties = form_data.get('involved_parties', [])
    prop = form_data.get('property_items', [])
    notif = form_data.get('notifications', {})

    for section in doc.sections:
        section.left_margin = section.right_margin = Inches(0.5)
        section.top_margin = section.bottom_margin = Inches(0.5)

    case_no = form_data.get('case_number') or '-'
    dept = officer.get('department_name') or '(department not set)'
    reported_dt = f"{inc.get('reported_date') or inc.get('date', '')} {inc.get('reported_time') or inc.get('time', '')}".strip()
    secure_dt = f"{inc.get('date', '')} {inc.get('time', '')}".strip()

    # 1. Header grid — band | agency/ORI/location | title + gang/premise/beat |
    #    stacked Case# / Date Reported / Last Known Secure / At Found column
    t = doc.add_table(rows=4, cols=6)
    t.style = 'Table Grid'
    _band_col(t, 'INCIDENT DATA')
    _grid_cell(t.rows[0].cells[1], 'Agency Name', dept)
    title_cell = t.rows[0].cells[2].merge(t.rows[1].cells[4])
    title_cell.text = ''
    p = title_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('INCIDENT/INVESTIGATION')
    run.add_break()
    run.add_text('REPORT')
    run.bold = True
    run.font.size = Pt(11)
    _grid_cell(t.rows[0].cells[5], 'Case#', case_no)
    _grid_cell(t.rows[1].cells[1], 'ORI', officer.get('ori') or '(ORI not set)')
    _grid_cell(t.rows[1].cells[5], 'Date / Time Reported', reported_dt)
    _grid_cell(t.rows[2].cells[1].merge(t.rows[3].cells[1]),
               'Location of Incident', inc.get('location') or '-')
    _grid_cell(t.rows[2].cells[2].merge(t.rows[3].cells[2]), 'Gang Relat', 'NO')
    _grid_cell(t.rows[2].cells[3].merge(t.rows[3].cells[3]), 'Premise Type',
               inc.get('premise_type') or 'Hotel/motel/etc.')
    _grid_cell(t.rows[2].cells[4].merge(t.rows[3].cells[4]), 'Beat/Tract', 'D')
    _grid_cell(t.rows[2].cells[5], 'Last Known Secure', secure_dt)
    _grid_cell(t.rows[3].cells[5], 'At Found', secure_dt)

    # 2. Crime incidents ×3 — main row + Entry/Exit/Security sub-row each
    categories = inc.get('categories', []) or ['General Information / Incident']
    weapon = (notif.get('weapon_detail') if notif.get('weapon_involved') else 'None') or 'None'
    t_cr = doc.add_table(rows=6, cols=5)
    t_cr.style = 'Table Grid'
    _band_col(t_cr, '')
    for idx in range(3):
        crime_name = categories[idx] if idx < len(categories) else ''
        filled = bool(crime_name)
        r0, r1 = idx * 2, idx * 2 + 1
        _grid_cell(t_cr.rows[r0].cells[1].merge(t_cr.rows[r1].cells[1]),
                   f"#{idx + 1} Crime Incident(s)  (Com)", crime_name)
        _grid_cell(t_cr.rows[r0].cells[2].merge(t_cr.rows[r0].cells[3]),
                   'Weapon / Tools', weapon if filled else '')
        _grid_cell(t_cr.rows[r0].cells[4], 'Activity', 'N' if filled else '')
        _grid_cell(t_cr.rows[r1].cells[2], 'Entry', 'None' if filled else '')
        _grid_cell(t_cr.rows[r1].cells[3], 'Exit', 'None' if filled else '')
        _grid_cell(t_cr.rows[r1].cells[4], 'Security', 'None' if filled else '')

    # 3. MO
    _grid(doc, [[('MO (Modus Operandi)', facts.get('how') or 'N/A')]], band='MO')

    # 4. Victim block (V1) — template column set
    victims = [p_ for p_ in parties if p_.get('role') == 'victim']
    others = [p_ for p_ in parties if p_.get('role') != 'victim']
    v = victims[0] if victims else {}
    veh = next((p_ for p_ in prop if p_.get('type') == 'vehicle'), {})
    resident = 'Resident' if v.get('address') else 'Non-Resident'
    _grid(doc, [
        [('# of Victims', str(max(1, len(victims)))),
         ('Type', 'INDIVIDUAL (NON LE)'),
         ('Injury', form_data.get('injuries', {}).get('description') or 'None'),
         ('Domestic', 'N')],
    ], band='VICTIM')
    _grid(doc, [
        [('V1 Victim/Business Name (Last, First, Middle)', v.get('full_name') or '-'),
         ('Victim of Crime #', '1'),
         ('DOB / Age', v.get('dob') or '-'),
         ('Race', v.get('race') or 'U'),
         ('Sex', v.get('sex') or 'U'),
         ('Relationship To Offender', 'INR'),
         ('Resident Status', resident),
         ('Military Branch/Status', '-')],
    ], band='')
    _grid(doc, [
        [('Home Address', v.get('address') or '-'),
         ('Email', v.get('email') or '-'),
         ('Home Phone', v.get('phone') or '-')],
        [('Employer Name/Address', '-'),
         ('Business Phone', '-'),
         ('Mobile Phone', v.get('phone') or '-')],
    ], band='')
    _grid(doc, [
        [('VYR', veh.get('year') or '-'),
         ('Make', veh.get('make') or '-'),
         ('Model', veh.get('model') or '-'),
         ('Style', '-'),
         ('Color', veh.get('color') or '-'),
         ('Lic/Lis', '-'),
         ('VIN', veh.get('serial_or_tag') or '-')],
    ], band='')

    # CODES legend
    _legend_row(doc, 'CODES:   V = Victim (Denote V2, V3)     WI = Witness     '
                     'IO = Involved Other     RP = Reporting Person (if other than victim)',
                band='')

    # 5. Others involved (two blocks, template-style)
    for idx in range(2):
        o = others[idx] if idx < len(others) else {}
        role_code = 'WI' if o.get('role') == 'witness' else 'IO'
        _grid(doc, [
            [('Type:', 'INDIVIDUAL (NON LE)' if o else ''),
             ('Injury:', '')],
        ], band='OTHERS INVOLVED' if idx == 0 else '')
        _grid(doc, [
            [('Code', role_code if o else ''),
             ('Name (Last, First, Middle)', o.get('full_name') or ('-' if o else '')),
             ('Victim of Crime #', ''),
             ('DOB / Age', o.get('dob') or ('-' if o else '')),
             ('Race', (o.get('race') or 'U') if o else ''),
             ('Sex', (o.get('sex') or 'U') if o else ''),
             ('Relationship To Offender', 'None' if o else ''),
             ('Resident Status', ('Resident' if o.get('address') else 'Non-Resident') if o else '')],
        ], band='')
        _grid(doc, [
            [('Home Address', o.get('address') or ('-' if o else '')),
             ('Email', o.get('email') or ('-' if o else '')),
             ('Home Phone', o.get('phone') or ('-' if o else ''))],
            [('Employer Name/Address', '-' if o else ''),
             ('Business Phone', '-' if o else ''),
             ('Mobile Phone', o.get('phone') or ('-' if o else ''))],
        ], band='')

    # 6. Property section — legend + column headers + full block of rows
    _legend_row(doc, '1 = None   2 = Burned   3 = Counterfeit / Forged   4 = Damaged / Vandalized   '
                     '5 = Recovered   6 = Seized   7 = Stolen   8 = Unknown', band='')
    headers = ['VI#', 'Code', 'Status', 'Value', 'OJ', 'QTY',
               'Property Description', 'Make/Model', 'Serial Number']
    prop_items = [p_ for p_ in prop if p_.get('type') != 'vehicle']
    t_prop = doc.add_table(rows=1 + max(8, len(prop_items)), cols=1 + len(headers))
    t_prop.style = 'Table Grid'
    _band_col(t_prop, 'PROPERTY')
    for c, h in enumerate(headers):
        cell = t_prop.rows[0].cells[c + 1]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(6)
    stat_map = {'missing': '7', 'stolen': '7', 'damaged': '4', 'recovered': '5', 'seized': '6'}
    for r in range(1, len(t_prop.rows)):
        p_item = prop_items[r - 1] if r - 1 < len(prop_items) else {}
        values = [
            'V1' if p_item else '', stat_map.get(p_item.get('status'), '1') if p_item else '',
            p_item.get('status') or '', f"${p_item.get('value')}" if p_item.get('value') else '',
            'N' if p_item else '', '1' if p_item else '', p_item.get('type', ''),
            f"{p_item.get('make') or ''} {p_item.get('model') or ''}".strip(),
            p_item.get('serial_or_tag') or '',
        ]
        for c, val in enumerate(values):
            cell = t_prop.rows[r].cells[c + 1]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(8)

    # 7. Officer / case status footer — Status band on the bottom row
    _grid(doc, [
        [('Officer / ID#', f"{officer.get('full_name') or '-'} ({officer.get('badge_number') or '-'})"),
         ('Invest ID# / Name', f"{officer.get('badge_number') or '-'} - {officer.get('full_name') or '-'}"),
         ('Supervisor', '_______________________')],
    ], band='')
    _grid(doc, [
        [('Complainant Signature', '_______________________'),
         ('Case Status', 'Closed By Investigation'),
         ('Case Disposition / Date', f"8 / {reported_dt.split(' ')[0] if reported_dt else '-'}")],
    ], band='Status')

    # 8. Page 2 — drugs / assisting officers / hate-bias (template page 2)
    doc.add_page_break()
    _title(doc, 'INCIDENT/INVESTIGATION REPORT')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{dept}    Case # {case_no}")
    run.font.size = Pt(8)
    _legend_row(doc, 'Status Codes:  1 = None   2 = Burned   3 = Counterfeit / Forged   '
                     '4 = Damaged / Vandalized   5 = Recovered   6 = Seized   7 = Stolen   8 = Unknown',
                band='')
    drug_headers = ['IBR', 'Status', 'Quantity', 'Type Measure', 'Suspected Type']
    drugs = form_data.get('drugs', [])
    t_drugs = doc.add_table(rows=1 + max(6, len(drugs)), cols=1 + len(drug_headers))
    t_drugs.style = 'Table Grid'
    _band_col(t_drugs, 'DRUGS')
    for c, h in enumerate(drug_headers):
        cell = t_drugs.rows[0].cells[c + 1]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(6)
    for r in range(1, len(t_drugs.rows)):
        d_item = drugs[r - 1] if r - 1 < len(drugs) else {}
        for c, key in enumerate(['ibr', 'status', 'quantity', 'type_measure', 'suspected_type']):
            cell = t_drugs.rows[r].cells[c + 1]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(d_item.get(key, '')))
            run.font.size = Pt(8)
    assisting = form_data.get('assisting_officers', [])
    if isinstance(assisting, list):
        assisting = ',  '.join(a for a in assisting if a)
    _grid(doc, [
        [('Assisting Officers', assisting or '')],
        [('Suspect Hate / Bias Motivated:', form_data.get('hate_bias', '') or '')],
    ], band='')
    doc.add_paragraph()

    # 9. Narrative continuation — template style: title, "Narr. (cont.) OCA:",
    #    agency line, then the boxed NARRATIVE section.
    _title(doc, 'INCIDENT/INVESTIGATION REPORT')
    t_nl = doc.add_table(rows=1, cols=2)
    left, right = t_nl.rows[0].cells
    left.text = ''
    run = left.paragraphs[0].add_run(f"Narr. (cont.)  OCA: {case_no}")
    run.font.size = Pt(8)
    right.text = ''
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(dept)
    run.italic = True
    run.font.size = Pt(8)
    _legend_row(doc, 'N A R R A T I V E')
    _narrative(doc, narrative)
    _signature(doc, officer)


# ── Search warrant — AO 93 layout ────────────────────────────────────
def _search_warrant(doc, form_data, narrative, officer, doc_meta=None):
    court = form_data.get('court', {})
    place = form_data.get('place_to_search', {})
    execution = form_data.get('execution', {})
    place_lines = [line for line in [place.get('description'), place.get('address')] if line]
    federal = _is_federal(court, officer)
    if federal:
        locality = (court.get('district') or officer.get('agency_judicial_district')
                    or 'the __________ District of __________')
        officer_kind = 'a federal law enforcement officer'
        judge_title = 'United States Magistrate Judge'
    else:
        county = officer.get('agency_county')
        locality = ', '.join(x for x in [f"{county} County" if county else '',
                                         officer.get('agency_state')] if x) or '____________________'
        officer_kind = 'a law enforcement officer'
        judge_title = officer.get('agency_judge_title') or 'Judge'

    if federal:
        _form_stamp(doc, 'AO 93 (Rev. 11/13) Search and Seizure Warrant')
    _draft_banner(doc, doc_meta)
    _court_head(doc, court, officer)
    _caption(doc, ['In the Matter of the Search of'] + (place_lines or ['____________________']),
             form_data.get('case_number'))
    _title(doc, 'SEARCH AND SEIZURE WARRANT')
    doc.add_paragraph('To:\tAny authorized law enforcement officer')

    doc.add_paragraph(
        f'An application by {officer_kind} or an attorney for the government '
        f'requests the search of the following person or property located in {locality} '
        '(identify the person or describe the property to be searched and give its location):')
    doc.add_paragraph('See Attachment A.')

    doc.add_paragraph(
        'I find that the affidavit(s), or any recorded testimony, establish probable cause to '
        'search and seize the person or property described above, and that such search will '
        'reveal (identify the person or describe the property to be seized):')
    doc.add_paragraph('See Attachment B.')

    p = doc.add_paragraph()
    p.add_run('YOU ARE COMMANDED').bold = True
    p.add_run(' to execute this warrant on or before '
              f"{execution.get('execute_by_date') or '____________'} (not to exceed 14 days)")
    anytime = execution.get('time_window') == 'anytime'
    _checkbox_line(doc, [
        ('daytime', 'in the daytime 6:00 a.m. to 10:00 p.m.'),
        ('anytime', 'at any time in the day or night because good cause has been established.'),
    ], 'anytime' if anytime else 'daytime')

    doc.add_paragraph(
        'Unless delayed notice is authorized below, you must give a copy of the warrant and a '
        'receipt for the property taken to the person from whom, or from whose premises, the '
        'property was taken, or leave the copy and receipt at the place where the property was taken.')
    doc.add_paragraph(
        'The officer executing this warrant, or an officer present during the execution of the '
        'warrant, must prepare an inventory as required by law and promptly return this warrant '
        f"and inventory to {court.get('judge_name') or '____________________'} "
        f'({judge_title}).')

    _sig_pair(doc, 'Date and time issued:', "Judge's signature")
    _sig_pair(doc, 'City and state:', 'Printed name and title')

    # Page 2 — Return
    doc.add_page_break()
    if federal:
        _form_stamp(doc, 'AO 93 (Rev. 11/13) Search and Seizure Warrant (Page 2)')
    _title(doc, 'Return')
    _kv(doc, [
        ('Case No.', form_data.get('case_number', '')),
        ('Date and time warrant executed', ''),
        ('Copy of warrant and inventory left with', ''),
        ('Inventory made in the presence of', ''),
        ('Inventory of the property taken and name of any person(s) seized', ''),
    ])
    _title(doc, 'Certification')
    doc.add_paragraph(
        'I declare under penalty of perjury that this inventory is correct and was returned '
        'along with the original warrant to the designated judge.')
    _sig_pair(doc, 'Date:', "Executing officer's signature")
    _sig_pair(doc, '', 'Printed name and title')

    # Attachments + affidavit referenced by the face form
    doc.add_page_break()
    _title(doc, 'ATTACHMENT A — Property to be Searched')
    doc.add_paragraph(place.get('description', '-'))
    doc.add_paragraph(f"Location: {place.get('address', '-')}")
    doc.add_page_break()
    _title(doc, 'ATTACHMENT B — Items to be Seized')
    for i, it in enumerate(form_data.get('items_to_seize', [])):
        doc.add_paragraph(f"{chr(97 + i)}. {it}")
    doc.add_page_break()
    _title(doc, 'AFFIDAVIT — Statement of Probable Cause')
    _narrative(doc, narrative)
    _signature(doc, officer, doc_meta)


# ── Arrest warrant — AO 442 layout ───────────────────────────────────
_CHARGING_DOCS_ROW1 = [
    ('indictment', 'Indictment'),
    ('superseding_indictment', 'Superseding Indictment'),
    ('information', 'Information'),
    ('superseding_information', 'Superseding Information'),
    ('complaint', 'Complaint'),
]
_CHARGING_DOCS_ROW2 = [
    ('probation_violation', 'Probation Violation Petition'),
    ('supervised_release_violation', 'Supervised Release Violation Petition'),
    ('violation_notice', 'Violation Notice'),
    ('court_order', 'Order of the Court'),
]


def _arrest_warrant(doc, form_data, narrative, officer, doc_meta=None):
    court = form_data.get('court', {})
    defendant = form_data.get('defendant', {})
    offense = form_data.get('offense', {})
    ident = form_data.get('identifiers', {})
    name = defendant.get('full_name', '')
    federal = _is_federal(court, officer)
    if federal:
        plaintiff = 'United States of America'
        judge_phrase = 'a United States magistrate judge'
        _form_stamp(doc, 'AO 442 (Rev. 11/11) Arrest Warrant')
    else:
        state = officer.get('agency_state')
        plaintiff = f"STATE OF {state.upper()}" if state else 'THE STATE'
        judge_phrase = 'a judge of this Court'

    _draft_banner(doc, doc_meta)
    _court_head(doc, court, officer)
    _caption(doc, [plaintiff, 'v.', name or '____________________', 'Defendant'],
             form_data.get('case_number'))
    _title(doc, 'ARREST WARRANT')
    doc.add_paragraph('To:\tAny authorized law enforcement officer')

    p = doc.add_paragraph()
    p.add_run('YOU ARE COMMANDED').bold = True
    p.add_run(f' to arrest and bring before {judge_phrase} without unnecessary '
              f"delay (name of person to be arrested) {name or '____________________'}, "
              'who is accused of an offense or violation based on the following document filed '
              'with the court:')
    selected = form_data.get('charging_document', '')
    _checkbox_line(doc, _CHARGING_DOCS_ROW1, selected)
    _checkbox_line(doc, _CHARGING_DOCS_ROW2, selected)

    doc.add_paragraph('This offense is briefly described as follows:')
    doc.add_paragraph(
        f"{offense.get('code_section', '')}  {offense.get('brief_description', '')}".strip() or '-')

    _sig_pair(doc, 'Date:', "Issuing officer's signature")
    _sig_pair(doc, 'City and state:', 'Printed name and title')

    _title(doc, 'Return')
    doc.add_paragraph(
        'This warrant was received on (date) ______________, and the person was arrested on '
        '(date) ______________ at (city and state) ____________________.')
    _sig_pair(doc, 'Date:', "Arresting officer's signature")
    _sig_pair(doc, '', 'Printed name and title')

    # Page 2 — sealed personal identifiers
    doc.add_page_break()
    if federal:
        _form_stamp(doc, 'AO 442 (Rev. 11/11) Arrest Warrant (Page 2)')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('This second page contains personal identifiers provided for law-enforcement use '
              'only and therefore should not be filed in court with the executed warrant unless '
              'under seal.').bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('(Not for Public Disclosure)').italic = True

    associates = '; '.join(
        f"{a.get('name', '')} ({a.get('relation', '')}) {a.get('phone', '')}".strip()
        for a in ident.get('known_associates', []))
    _field(doc, 'Name of defendant/offender', name)
    _field(doc, 'Known aliases', ', '.join(ident.get('aliases', [])))
    _field(doc, 'Last known residence', ident.get('last_known_residence', ''))
    _field(doc, 'Prior addresses to which defendant/offender may still have ties',
           '; '.join(ident.get('prior_addresses', [])))
    _field(doc, 'Last known employment', ident.get('last_known_employment', ''))
    _field(doc, 'Last known telephone numbers', ', '.join(ident.get('phone_numbers', [])))
    _field(doc, 'Place of birth', ident.get('place_of_birth', ''))
    _field(doc, 'Date of birth', ident.get('date_of_birth', ''))
    _field(doc, 'Social Security number', ident.get('ssn', ''))
    _field(doc, 'Height', ident.get('height', ''))
    _field(doc, 'Weight', ident.get('weight', ''))
    _field(doc, 'Sex', ident.get('sex', ''))
    _field(doc, 'Race', ident.get('race', ''))
    _field(doc, 'Hair', ident.get('hair', ''))
    _field(doc, 'Eyes', ident.get('eyes', ''))
    _field(doc, 'Scars, tattoos, other distinguishing marks', ident.get('distinguishing_marks', ''))
    _field(doc, 'History of violence, weapons, drug use',
           ident.get('history_violence_weapons_drugs', ''))
    _field(doc, 'Known family, friends, and other associates (name, relation, address, phone number)',
           associates)
    _field(doc, 'FBI number', ident.get('fbi_number', ''))
    _field(doc, 'Complete description of auto', ident.get('vehicle_description', ''))
    _field(doc, 'Investigative agency and address', ident.get('investigative_agency', ''))

    if narrative and narrative.strip():
        doc.add_page_break()
        _title(doc, 'SUPPORTING AFFIDAVIT')
        _narrative(doc, narrative)
        _signature(doc, officer, doc_meta)


_BUILDERS = {
    'incident_report': _incident,
    'search_warrant': _search_warrant,
    'arrest_warrant': _arrest_warrant,
}


def render_docx(doc_type, form_data, narrative, officer, doc_meta=None) -> io.BytesIO:
    builder = _BUILDERS.get(doc_type)
    if builder is None:
        raise ValueError(f'No DOCX template for doc_type {doc_type}')
    doc = Document()
    if doc_type == 'incident_report':
        builder(doc, form_data, narrative, officer)
    else:
        builder(doc, form_data, narrative, officer, doc_meta)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
